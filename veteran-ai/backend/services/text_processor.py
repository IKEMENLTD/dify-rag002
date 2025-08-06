import re
import pytesseract
import whisper
from PIL import Image
from pdf2image import convert_from_path
import PyPDF2
from docx import Document
from typing import Optional
import os
import tempfile
from pydub import AudioSegment

class TextProcessor:
    def __init__(self):
        # Load Whisper model for audio processing
        try:
            self.whisper_model = whisper.load_model("small")  # Can be changed to "base", "large"
        except Exception as e:
            print(f"Whisper model loading error: {e}")
            self.whisper_model = None
    
    async def process_chat_message(self, message: str) -> str:
        """Process and clean chat message text"""
        if not message:
            return ""
        
        # Remove URLs
        message = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '[URL]', message)
        
        # Remove email addresses
        message = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[EMAIL]', message)
        
        # Remove excessive whitespace
        message = re.sub(r'\s+', ' ', message).strip()
        
        # Remove Slack/Discord formatting
        message = re.sub(r'<@[A-Z0-9]+>', '@user', message)  # User mentions
        message = re.sub(r'<#[A-Z0-9]+\|[^>]+>', '#channel', message)  # Channel mentions
        message = re.sub(r'<[^>]+>', '', message)  # Other formatting
        
        return message
    
    async def process_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            # Open image
            image = Image.open(file_path)
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(image, lang='jpn+eng')
            
            # Clean extracted text
            text = self._clean_extracted_text(text)
            
            return text
            
        except Exception as e:
            print(f"Image processing error: {e}")
            return ""
    
    async def process_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            text = ""
            
            # Try text extraction first
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += page_text + "\n"
            
            # If no text found, try OCR on images
            if not text.strip():
                try:
                    images = convert_from_path(file_path)
                    for image in images:
                        image_text = pytesseract.image_to_string(image, lang='jpn+eng')
                        text += image_text + "\n"
                except Exception as ocr_error:
                    print(f"PDF OCR error: {ocr_error}")
            
            # Clean extracted text
            text = self._clean_extracted_text(text)
            
            return text
            
        except Exception as e:
            print(f"PDF processing error: {e}")
            return ""
    
    async def process_audio(self, file_path: str) -> str:
        """Convert audio to text using Whisper"""
        try:
            if not self.whisper_model:
                return ""
            
            # Convert to wav if necessary
            temp_wav_path = None
            
            if not file_path.lower().endswith('.wav'):
                # Convert to wav
                audio = AudioSegment.from_file(file_path)
                temp_wav_path = tempfile.NamedTemporaryFile(suffix='.wav', delete=False).name
                audio.export(temp_wav_path, format="wav")
                process_path = temp_wav_path
            else:
                process_path = file_path
            
            # Transcribe using Whisper
            result = self.whisper_model.transcribe(process_path, language='ja')
            text = result["text"]
            
            # Clean up temporary file
            if temp_wav_path and os.path.exists(temp_wav_path):
                os.unlink(temp_wav_path)
            
            # Clean extracted text
            text = self._clean_extracted_text(text)
            
            return text
            
        except Exception as e:
            print(f"Audio processing error: {e}")
            return ""
    
    async def process_text_file(self, file_path: str) -> str:
        """Process plain text files"""
        try:
            # Detect file type and process accordingly
            if file_path.lower().endswith('.docx'):
                return await self._process_docx(file_path)
            elif file_path.lower().endswith(('.txt', '.md', '.csv')):
                return await self._process_plain_text(file_path)
            else:
                return await self._process_plain_text(file_path)
                
        except Exception as e:
            print(f"Text file processing error: {e}")
            return ""
    
    async def _process_docx(self, file_path: str) -> str:
        """Process Word document"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_extracted_text(text)
            
        except Exception as e:
            print(f"DOCX processing error: {e}")
            return ""
    
    async def _process_plain_text(self, file_path: str) -> str:
        """Process plain text files with encoding detection"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'shift_jis', 'cp932', 'euc-jp', 'iso-2022-jp']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        return self._clean_extracted_text(text)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
                return self._clean_extracted_text(text)
                
        except Exception as e:
            print(f"Plain text processing error: {e}")
            return ""
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove non-printable characters except Japanese
        text = re.sub(r'[^\w\s\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FAF\u002D\u002E\u003A\u003B\u003F\u0021\u002C\u0028\u0029\u005B\u005D\u007B\u007D]', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n+', '\n', text)
        
        # Trim
        text = text.strip()
        
        return text
    
    def extract_metadata(self, text: str) -> dict:
        """Extract metadata from text"""
        metadata = {
            "length": len(text),
            "word_count": len(text.split()),
            "contains_japanese": bool(re.search(r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FAF]', text)),
            "contains_english": bool(re.search(r'[a-zA-Z]', text)),
            "contains_numbers": bool(re.search(r'\d', text)),
        }
        
        # Extract potential dates
        date_patterns = [
            r'\d{4}[-/]\d{1,2}[-/]\d{1,2}',
            r'\d{1,2}[-/]\d{1,2}[-/]\d{4}',
            r'\d{4}年\d{1,2}月\d{1,2}日'
        ]
        
        dates = []
        for pattern in date_patterns:
            dates.extend(re.findall(pattern, text))
        
        if dates:
            metadata["dates_found"] = list(set(dates))
        
        return metadata

# Global instance
text_processor = TextProcessor()